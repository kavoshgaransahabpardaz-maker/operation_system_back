"""
Tabular file ingestion — parse XLS/XLSX/CSV/XML directly without OCR.
PURE PYTHON — no LLM, no network.
"""
import csv
import io

_TABULAR_EXTENSIONS = {".xls", ".xlsx", ".csv", ".xml", ".docx"}
_TABULAR_CONTENT_TYPES = {
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/csv",
    "application/xml",
    "text/xml",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def is_tabular(filename: str, content_type: str) -> bool:
    """Return True if the file should use the tabular path instead of OCR."""
    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()
    return ext in _TABULAR_EXTENSIONS or content_type.lower().split(";")[0].strip() in _TABULAR_CONTENT_TYPES


def parse_tabular(file_bytes: bytes, filename: str, content_type: str) -> str:
    """
    Parse a tabular file and return its content as plain text suitable for field extraction.

    - XLS/XLSX: converted to CSV-like text via openpyxl
    - CSV: decoded and returned as-is
    - XML: pretty-printed element tree via lxml

    Raises ValueError if the format is not recognised or parsing fails.
    """
    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()
    ct = content_type.lower().split(";")[0].strip()

    # --- XLS / XLSX ---
    if ext in {".xls", ".xlsx"} or ct in {
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }:
        return _parse_excel(file_bytes, filename)

    # --- CSV ---
    if ext == ".csv" or ct == "text/csv":
        return _parse_csv(file_bytes)

    # --- XML ---
    if ext == ".xml" or ct in {"application/xml", "text/xml"}:
        return _parse_xml(file_bytes)

    # --- DOCX ---
    if ext == ".docx" or ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_bytes)

    raise ValueError(
        f"Unsupported tabular format: extension={ext!r}, content_type={content_type!r}"
    )


def _parse_excel(file_bytes: bytes, filename: str) -> str:
    try:
        import openpyxl  # type: ignore
    except ImportError as exc:
        raise ValueError("openpyxl is required for XLS/XLSX parsing") from exc

    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError(f"Failed to open workbook '{filename}': {exc}") from exc

    lines: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        buf = io.StringIO()
        writer = csv.writer(buf)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(["" if v is None else str(v) for v in row])
        sheet_text = buf.getvalue().strip()
        if sheet_text:
            lines.append(sheet_text)

    wb.close()
    return "\n".join(lines)


def _parse_csv(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("CSV file could not be decoded with utf-8 or latin-1 encoding")


def _parse_xml(file_bytes: bytes) -> str:
    try:
        from lxml import etree  # type: ignore
    except ImportError as exc:
        raise ValueError("lxml is required for XML parsing") from exc

    try:
        root = etree.fromstring(file_bytes)  # noqa: S320  (not user-controlled in prod)
    except etree.XMLSyntaxError as exc:
        raise ValueError(f"XML parse error: {exc}") from exc

    lines: list[str] = []
    _xml_to_lines(root, lines, indent=0)
    return "\n".join(lines)


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as exc:
        raise ValueError("python-docx is required for DOCX parsing") from exc

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"Failed to open DOCX: {exc}") from exc

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _xml_to_lines(element, lines: list[str], indent: int) -> None:
    """Recursively convert an lxml element tree into indented text."""
    from lxml import etree  # type: ignore

    tag = etree.QName(element.tag).localname if isinstance(element.tag, str) else str(element.tag)
    attrs = " ".join(f'{k}="{v}"' for k, v in element.attrib.items())
    header = f"{'  ' * indent}<{tag}{' ' + attrs if attrs else ''}>"
    text = (element.text or "").strip()
    if text:
        header += f" {text}"
    lines.append(header)
    for child in element:
        _xml_to_lines(child, lines, indent + 1)
